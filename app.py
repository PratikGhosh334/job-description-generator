from flask import Flask, request, render_template, jsonify, send_file
import google.generativeai as genai
import os
from dotenv import load_dotenv
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO
from docx import Document

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

# Configure the Gemini API with your API key
genai.configure(api_key=os.environ["GOOGLE_API_KEY"])  # Ensure your API key is set as an environment variable

def call_gemini_api(prompt):
    # Call the Gemini API to generate content
    model = genai.GenerativeModel("gemini-1.5-flash")  # Use the appropriate model
    response = model.generate_content(prompt)
    return response.text  # Return the generated text

def clean_response(text):
    # Remove double stars (**) from headings
    text = text.replace('**', '')
    
    # Replace single stars (*) with dashes (-)
    text = text.replace('*', '-')
    
    # Remove extra whitespace
    text = '\n'.join([line.strip() for line in text.split('\n')])
    
    return text

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate_job_description():
    job_title = request.form['job_title']
    skills = request.form['skills']
    culture = request.form['culture']
    details = request.form['details']

    # Generate the job description using LangChain and Gemini API
    prompt = f"Generate a job description for a {job_title} role. The required skills are {skills}. The company culture is {culture}. Additional details: {details}."
    
    # Call the Gemini API
    job_description = call_gemini_api(prompt)

    # Clean up the response
    cleaned_description = clean_response(job_description)

    return jsonify({'job_description': cleaned_description})

@app.route('/export_pdf', methods=['POST'])
def export_pdf():
    job_description = request.form['job_description']
    
    # Create a PDF in memory
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter

    # Draw the job description on the PDF
    p.drawString(100, height - 100, "Job Description")
    p.drawString(100, height - 120, job_description)

    p.showPage()
    p.save()

    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name='job_description.pdf', mimetype='application/pdf')

@app.route('/export_docx', methods=['POST'])
def export_docx():
    job_description = request.json['job_description']
    
    # Create a DOCX document in memory
    doc = Document()
    doc.add_heading('Job Description', level=1)
    doc.add_paragraph(job_description)

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name='job_description.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
